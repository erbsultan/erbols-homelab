from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


OUT_DIR = Path("/Users/erbol/erbols-homelab/resume")
DOCX_PATH = OUT_DIR / "Erbol_Sultanbayev_DevOps_Resume_Updated.docx"
PDF_PATH = OUT_DIR / "Erbol_Sultanbayev_DevOps_Resume_Updated.pdf"

CONTACT = [
    "+998 50 515 02 76",
    "m.erbol.sultanbayev@gmail.com",
    "Telegram: @erbol_sultanbayev",
    "LinkedIn: erbol-sultanbayev",
    "GitHub: github.com/erbsultan",
    "Tashkent, Uzbekistan",
]

SUMMARY = (
    "Junior DevOps engineer in training building a real homelab on a hardened Ubuntu VPS. "
    "Hands-on with Linux server administration, Docker Compose, nginx, CI/CD with GitHub Actions, "
    "monitoring with Prometheus/Grafana, log collection with Loki/Alloy, and Telegram alerting. "
    "Focused on practical infrastructure, observability, and reliable deployment workflows."
)

SKILLS = [
    ("Linux & server administration", "Ubuntu 24.04, systemd basics, SSH key-only access, ufw, fail2ban, unattended upgrades"),
    ("Containers", "Docker, Docker Compose, container networking, persistent volumes"),
    ("Web & TLS", "nginx static hosting and reverse proxy, Let's Encrypt, certbot, HTTP to HTTPS redirects"),
    ("CI/CD", "Git, GitHub, GitHub Actions, rsync deploys, SSH deploy keys, workflow secrets"),
    ("Observability", "Grafana, Prometheus, node_exporter, Loki, Grafana Alloy, LogQL basics, Grafana Alerting"),
    ("Automation & scripting", "Bash basics, Python for utility scripts, Telegram Bot API notifications"),
    ("Networking", "DNS records, TCP/IP basics, HTTP/HTTPS, localhost binding, private service exposure"),
]

PROJECTS = [
    (
        "erbols-homelab - live VPS infrastructure",
        "2026",
        [
            "Built and documented a public homelab on an Ubuntu 24.04 VPS hosted on Vultr.",
            "Serves erbsultan.uz through nginx with Let's Encrypt TLS, DNS records, HTTP to HTTPS redirect, and a hardened SSH/firewall baseline.",
            "Implemented GitHub Actions deployment with rsync over SSH and Telegram status notifications for successful or failed deploys.",
        ],
    ),
    (
        "Observability and alerting stack",
        "2026",
        [
            "Deployed Grafana, Prometheus, node_exporter, Loki, and Alloy with Docker Compose.",
            "Configured Prometheus scraping for VPS metrics and Loki log collection for nginx access/error logs.",
            "Provisioned Grafana data sources, alert rules, notification policy, and Telegram contact point from repository files.",
            "Added scheduled GitHub Actions smoke checks for erbsultan.uz and grafana.erbsultan.uz with Telegram failure notifications.",
        ],
    ),
]

EDUCATION = [
    "School 21 - DevOps / systems track, 2024 - Present",
    "EPAM Learn - Cloud & DevOps, online",
    "KodeKloud DevOps Mastery - self-paced practice",
]

LANGUAGES = "Russian - Native | Kazakh - Native | Uzbek - Native | English - Intermediate (B1)"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D9D9D9"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "4")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)
        borders.append(tag)
    tc_pr.append(borders)


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(31, 78, 121)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Inches(0.22)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    p.add_run(text).font.size = Pt(9.2)


def build_docx():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)

    normal = doc.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    normal.paragraph_format.space_after = Pt(4)
    normal.paragraph_format.line_spacing = 1.05

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run("Erbol Sultanbayev")
    r.bold = True
    r.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.space_after = Pt(2)
    r = subtitle.add_run("Junior DevOps Engineer | Linux, CI/CD, Observability")
    r.font.size = Pt(10.5)
    r.font.color.rgb = RGBColor(80, 80, 80)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(8)
    r = contact.add_run(" | ".join(CONTACT))
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(80, 80, 80)

    add_heading(doc, "Summary")
    p = doc.add_paragraph(SUMMARY)
    p.paragraph_format.space_after = Pt(6)
    p.runs[0].font.size = Pt(9.4)

    add_heading(doc, "Hard Skills")
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    for category, value in SKILLS:
        cells = table.add_row().cells
        cells[0].width = Inches(1.75)
        cells[1].width = Inches(5.35)
        for cell in cells:
            set_cell_border(cell)
        set_cell_shading(cells[0], "F2F4F7")
        r0 = cells[0].paragraphs[0].add_run(category)
        r0.bold = True
        r0.font.size = Pt(8.8)
        r1 = cells[1].paragraphs[0].add_run(value)
        r1.font.size = Pt(8.8)

    add_heading(doc, "Projects")
    for name, date, bullets in PROJECTS:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(1)
        r = p.add_run(name)
        r.bold = True
        r.font.size = Pt(10)
        r2 = p.add_run(f"  {date}")
        r2.font.size = Pt(9)
        r2.font.color.rgb = RGBColor(100, 100, 100)
        for bullet in bullets:
            add_bullet(doc, bullet)

    add_heading(doc, "Education")
    for item in EDUCATION:
        add_bullet(doc, item)

    add_heading(doc, "Languages")
    p = doc.add_paragraph(LANGUAGES)
    p.runs[0].font.size = Pt(9.2)

    doc.save(DOCX_PATH)


def build_pdf():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    styles = getSampleStyleSheet()
    base = ParagraphStyle("Base", parent=styles["Normal"], fontName="Helvetica", fontSize=8.7, leading=10.7, spaceAfter=3)
    small = ParagraphStyle("Small", parent=base, fontSize=7.7, leading=9.2, textColor=colors.HexColor("#555555"), alignment=1)
    title = ParagraphStyle("Title", parent=base, fontName="Helvetica-Bold", fontSize=21, leading=24, alignment=1, spaceAfter=1)
    subtitle = ParagraphStyle("Subtitle", parent=small, fontSize=9.5, leading=11, spaceAfter=4)
    heading = ParagraphStyle("Heading", parent=base, fontName="Helvetica-Bold", fontSize=9.2, leading=11, textColor=colors.HexColor("#1F4E79"), spaceBefore=6, spaceAfter=3)
    bullet = ParagraphStyle("Bullet", parent=base, leftIndent=12, firstLineIndent=-7, spaceAfter=1.5)
    project_title = ParagraphStyle("ProjectTitle", parent=base, fontName="Helvetica-Bold", fontSize=9.5, leading=11, spaceBefore=3, spaceAfter=1)

    doc = SimpleDocTemplate(
        str(PDF_PATH),
        pagesize=letter,
        leftMargin=0.55 * inch,
        rightMargin=0.55 * inch,
        topMargin=0.45 * inch,
        bottomMargin=0.45 * inch,
    )

    story = [
        Paragraph("Erbol Sultanbayev", title),
        Paragraph("Junior DevOps Engineer | Linux, CI/CD, Observability", subtitle),
        Paragraph(" | ".join(CONTACT), small),
        Spacer(1, 4),
        Paragraph("SUMMARY", heading),
        Paragraph(SUMMARY, base),
        Paragraph("HARD SKILLS", heading),
    ]

    data = [[Paragraph(f"<b>{category}</b>", base), Paragraph(value, base)] for category, value in SKILLS]
    table = Table(data, colWidths=[1.65 * inch, 5.65 * inch], hAlign="LEFT")
    table.setStyle(TableStyle([
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#F2F4F7")),
        ("BOX", (0, 0), (-1, -1), 0.25, colors.HexColor("#D9D9D9")),
        ("INNERGRID", (0, 0), (-1, -1), 0.25, colors.HexColor("#D9D9D9")),
        ("LEFTPADDING", (0, 0), (-1, -1), 5),
        ("RIGHTPADDING", (0, 0), (-1, -1), 5),
        ("TOPPADDING", (0, 0), (-1, -1), 3),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 3),
    ]))
    story.append(table)

    story.append(Paragraph("PROJECTS", heading))
    for name, date, bullets in PROJECTS:
        story.append(Paragraph(f"{name} <font color='#666666'>{date}</font>", project_title))
        for item in bullets:
            story.append(Paragraph(f"- {item}", bullet))

    story.append(Paragraph("EDUCATION", heading))
    for item in EDUCATION:
        story.append(Paragraph(f"- {item}", bullet))

    story.append(Paragraph("LANGUAGES", heading))
    story.append(Paragraph(LANGUAGES, base))
    def white_page(canvas, _doc):
        canvas.saveState()
        canvas.setFillColor(colors.white)
        canvas.rect(0, 0, letter[0], letter[1], stroke=0, fill=1)
        canvas.restoreState()

    doc.build(story, onFirstPage=white_page, onLaterPages=white_page)


if __name__ == "__main__":
    build_docx()
    build_pdf()
    print(DOCX_PATH)
    print(PDF_PATH)
